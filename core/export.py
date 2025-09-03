import json
from io import BytesIO
from docx import Document
from .story import StrategyStory

def export_json(story: StrategyStory) -> str:
    return json.dumps(story.to_dict(), indent=2)

def export_docx_bytes(story: StrategyStory) -> bytes:
    doc = Document()
    doc.add_heading('Strategy Story — Steps 1–3', 0)

    doc.add_heading('Prompt', level=1)
    doc.add_paragraph(story.prompt or "—")

    c = story.clarifications
    doc.add_heading('Clarifications (Step 1)', level=1)
    for k in ["focus_area", "purpose", "industry", "geography", "time_horizon"]:
        doc.add_paragraph(f"{k.replace('_',' ').title()}: {c.get(k) or '—'}")
    outcomes = c.get("decision_outcomes", [])
    doc.add_paragraph("Decision Outcomes: " + (", ".join(outcomes) if outcomes else "—"))

    doc.add_heading('Assessments (Step 2)', level=1)
    sel = story.assessments.get("selected", [])
    if sel:
        for a in sel:
            doc.add_paragraph(a, style='List Bullet')
    else:
        doc.add_paragraph("—")

    doc.add_heading('Sub-Assessments (Step 3)', level=1)
    subs = story.sub_assessments.get("selected", {})
    if subs:
        for a, items in subs.items():
            doc.add_heading(a, level=2)
            if items:
                for it in items:
                    doc.add_paragraph(it, style='List Bullet')
            else:
                doc.add_paragraph("—")
    else:
        doc.add_paragraph("—")

    doc.add_paragraph()
    doc.add_paragraph(f"Updated at (UTC): {story.updated_at or '—'}")

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
